#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys
from argparse import ArgumentParser
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

def process(document, line, style, codestyle):
    if style==-1:
        paragraph = document.add_paragraph().add_run(unicode(line, "utf-8"), style=codestyle)
    else:
        document.add_heading(unicode(line, "utf-8"), style)

parser = ArgumentParser()

parser.add_argument('-f', '--file', default='result.docx', help='file to save')
parser.add_argument('-s', '--style', default=-1, type=int, help='is heading')
parser.add_argument('-t', '--translate', default='0', type=int, help='translate to stdout')
parser.add_argument('-c', '--codestyle', default='CodeStyle', help='is code')
args = parser.parse_args()

filename = args.file
style = args.style
translate = args.translate
codestyle = args.codestyle

if not os.path.exists(filename):
    document = Document()
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CodeStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(10)
    obj_font.name = 'Courier New'
    obj_charstyle = obj_styles.add_style('TextStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'
else:
    document = Document(filename)

while True:
    line = sys.stdin.readline()
    if not line:
        break
    if translate:
        print(unicode(line, "utf-8").rstrip('\n'))
    process(document, line.rstrip(), style, codestyle)
"""
document.add_heading('Document Title', 1)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run(u' перенос ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

# document.add_page_break()
"""
document.save(filename)

